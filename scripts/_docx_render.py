"""Shared Word (.docx) renderer for NEET subject categorization documents.

Produces the same structure as the PDF renderer: title, sources, distribution
table with class-level split, per-question table, notes section. Designed so
the chemistry/biology/physics build scripts can reuse the same data dicts.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BRAND = RGBColor(0x0D, 0x6E, 0xFD)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT_HEX = "F8FAFC"
ROW_TOTAL_HEX = "EEF2FF"
BORDER_HEX = "CBD5E1"

# Convert a tiny subset of inline HTML-ish tags from the PDF data into
# runs with bold/italic. Supports <b>...</b> only (sufficient for our data).
_BOLD_RE = re.compile(r"<b>(.*?)</b>", re.DOTALL)


def _add_runs(paragraph, text: str, *, base_size=10.0, base_bold=False,
              color: RGBColor | None = None):
    """Append text to a paragraph, honouring <b>...</b> markup."""
    text = text.replace("&nbsp;", "\u00a0")
    cursor = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > cursor:
            r = paragraph.add_run(text[cursor:m.start()])
            r.font.size = Pt(base_size)
            r.bold = base_bold
            if color is not None:
                r.font.color.rgb = color
        r = paragraph.add_run(m.group(1))
        r.font.size = Pt(base_size)
        r.bold = True
        if color is not None:
            r.font.color.rgb = color
        cursor = m.end()
    if cursor < len(text):
        r = paragraph.add_run(text[cursor:])
        r.font.size = Pt(base_size)
        r.bold = base_bold
        if color is not None:
            r.font.color.rgb = color


def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_borders(cell, hex_color: str = BORDER_HEX, size: str = "4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), hex_color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_page_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def _add_header(paragraph, text: str, *, size: float, color: RGBColor, bold=True,
                space_before=0.0, space_after=4.0):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    tr_pr.append(th)


def _populate_cell(cell, text: str, *, size: float = 9.0, bold=False,
                    color: RGBColor | None = None,
                    bg: str | None = None, align=None):
    cell.text = ""  # clear default empty paragraph
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _add_runs(p, text, base_size=size, base_bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if bg else WD_ALIGN_VERTICAL.TOP
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_borders(cell)


def render(
    *,
    out_path: Path,
    subject: str,
    q_range: str,
    total_q: int,
    sources_blurb: str,
    distribution: Sequence[tuple[str, str, int, str]],
    class_split: Sequence[str] | None,
    rows: Sequence[tuple[int, str, str, str, str]],
    notes: Sequence[tuple[str, str]],
    dist_col_widths_cm: tuple[float, float, float, float] = (1.5, 9.5, 1.8, 5.2),
    rows_col_widths_cm: tuple[float, float, float, float, float] = (1.0, 1.2, 4.0, 5.0, 6.8),
):
    """Build the .docx. Argument shape mirrors `_pdf_render.render`."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_page_margins(doc)
    doc.core_properties.title = f"NEET-UG 2026 — {subject} Categorization"
    doc.core_properties.author = "NEETSurge"

    # --- Title and intro ---
    p_title = doc.add_paragraph()
    _add_header(
        p_title,
        f"NEET-UG 2026 — {subject} Categorization",
        size=20, color=INK, bold=True, space_after=4,
    )

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(4)
    _add_runs(
        p_intro,
        f"Each of the {total_q} {subject} questions ({q_range}) in the "
        "NEET-UG 2026 paper, mapped to its corresponding Unit / Subtopic in the "
        "official NEET-UG 2026 syllabus.",
        base_size=10.5,
    )

    p_src = doc.add_paragraph()
    p_src.paragraph_format.space_after = Pt(10)
    r = p_src.add_run("Sources: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    r2 = p_src.add_run(sources_blurb)
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MUTED

    # --- Distribution heading ---
    p_h = doc.add_paragraph()
    _add_header(
        p_h,
        "Distribution across the syllabus units",
        size=14, color=BRAND, space_before=6, space_after=6,
    )

    # Distribution table
    headers = ["Unit", "Syllabus Unit", "Q count", "Question Numbers"]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    for i, w in enumerate(dist_col_widths_cm):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)
    head = tbl.rows[0]
    _set_repeat_header(head)
    for i, h in enumerate(headers):
        _populate_cell(head.cells[i], h, size=10, bold=True, color=WHITE, bg="0D6EFD",
                       align=WD_ALIGN_PARAGRAPH.LEFT)
    head.height = Cm(0.7)
    head.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for idx, (unit, name, count, qs) in enumerate(distribution):
        row = tbl.add_row()
        for i, w in enumerate(dist_col_widths_cm):
            row.cells[i].width = Cm(w)
        bg = ROW_ALT_HEX if idx % 2 == 1 else None
        _populate_cell(row.cells[0], unit, size=9.5, bg=bg)
        _populate_cell(row.cells[1], name, size=9.5, bg=bg)
        _populate_cell(row.cells[2], str(count), size=9.5, bg=bg)
        _populate_cell(row.cells[3], qs, size=9.5, bg=bg)

    total_row = tbl.add_row()
    for i, w in enumerate(dist_col_widths_cm):
        total_row.cells[i].width = Cm(w)
    _populate_cell(total_row.cells[0], "<b>Total</b>", size=9.5, bg=ROW_TOTAL_HEX)
    _populate_cell(total_row.cells[1], "", size=9.5, bg=ROW_TOTAL_HEX)
    _populate_cell(total_row.cells[2], f"<b>{total_q}</b>", size=9.5, bg=ROW_TOTAL_HEX)
    _populate_cell(total_row.cells[3], "", size=9.5, bg=ROW_TOTAL_HEX)

    # Class split
    if class_split:
        p_cs_h = doc.add_paragraph()
        _add_header(p_cs_h, "Class-level split (NCERT)", size=11, color=INK,
                    space_before=8, space_after=4)
        for line in class_split:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("• ")
            r.font.size = Pt(10.5)
            _add_runs(p, line, base_size=10.5)

    # --- Per-question section ---
    doc.add_page_break()
    p_h2 = doc.add_paragraph()
    _add_header(p_h2, "Per-question categorization", size=14, color=BRAND,
                space_before=0, space_after=6)

    row_headers = [
        "Q#", "Ans", "Syllabus Unit", "Subtopic",
        "What the question tests (from the PDF explanation)",
    ]
    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.autofit = False
    for i, w in enumerate(rows_col_widths_cm):
        for cell in tbl2.columns[i].cells:
            cell.width = Cm(w)
    head2 = tbl2.rows[0]
    _set_repeat_header(head2)
    for i, h in enumerate(row_headers):
        _populate_cell(head2.cells[i], h, size=10, bold=True, color=WHITE, bg="0D6EFD",
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    for idx, (q, ans, unit, sub, test) in enumerate(rows):
        row = tbl2.add_row()
        for i, w in enumerate(rows_col_widths_cm):
            row.cells[i].width = Cm(w)
        bg = ROW_ALT_HEX if idx % 2 == 1 else None
        _populate_cell(row.cells[0], f"<b>{q}</b>", size=9.5, bg=bg)
        _populate_cell(row.cells[1], f"<b>{ans}</b>", size=9.5, bg=bg)
        _populate_cell(row.cells[2], unit, size=9.5, bg=bg)
        _populate_cell(row.cells[3], sub, size=9.5, bg=bg)
        _populate_cell(row.cells[4], test, size=9.5, bg=bg)

    # --- Notes section ---
    if notes:
        doc.add_page_break()
        p_h3 = doc.add_paragraph()
        _add_header(p_h3, "Notes on borderline / cross-listed items", size=14,
                    color=BRAND, space_before=0, space_after=6)
        intro = doc.add_paragraph()
        intro.paragraph_format.space_after = Pt(6)
        _add_runs(
            intro,
            "A few questions sit at the boundary of two syllabus units. "
            "These are the choices made and why:",
            base_size=10.5,
        )
        for label, body in notes:
            p_lab = doc.add_paragraph()
            _add_header(p_lab, label, size=11, color=INK, space_before=6,
                        space_after=2)
            p_body = doc.add_paragraph()
            p_body.paragraph_format.space_after = Pt(4)
            _add_runs(p_body, body, base_size=10.5)

    doc.save(str(out_path))
    print(f"Wrote {out_path}  ({out_path.stat().st_size // 1024} KB)")
